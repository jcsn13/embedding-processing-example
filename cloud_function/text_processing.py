"""
Copyright 2025 Google LLC

Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at

     https://www.apache.org/licenses/LICENSE-2.0

Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

import os
import logging
from typing import Optional, Dict, Any
import tempfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def extract_text_from_document(file_path: str) -> str:
    """
    Extract text from a document file using unstructured library.

    Args:
        file_path (str): Path to the document file

    Returns:
        str: Extracted text
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    logger.info(f"Detected file extension: '{file_extension}' for file: {file_path}")

    # If no extension is detected, try to determine file type
    if not file_extension:
        file_extension = detect_file_type(file_path)
        logger.info(f"Detected file type: '{file_extension}' using content analysis")

    try:
        # Use unstructured for document processing
        from unstructured.partition.auto import partition

        logger.info(f"Extracting text from {file_path} using unstructured")
        elements = partition(file_path)
        text = "\n\n".join([str(element) for element in elements])

        # Clean the extracted text
        text = clean_text(text)

        logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
        return text

    except ImportError:
        logger.warning(
            "unstructured library not available, falling back to specific extractors"
        )
        return fallback_extract_text(file_path)
    except Exception as e:
        logger.error(f"Error using unstructured: {str(e)}")
        logger.warning("Falling back to specific extractors")
        return fallback_extract_text(file_path)


def detect_file_type(file_path: str) -> str:
    """
    Detect file type based on content when extension is missing.

    Args:
        file_path (str): Path to the file

    Returns:
        str: Detected file extension including the dot (e.g., '.pdf')
    """
    try:
        # Read the first few bytes to check for file signatures
        with open(file_path, "rb") as f:
            header = f.read(8)

        # Check for PDF signature (%PDF-)
        if header.startswith(b"%PDF-"):
            logger.info(f"File signature indicates PDF format")
            return ".pdf"

        # Check for DOCX (ZIP format with specific content)
        if header.startswith(b"PK\x03\x04"):
            logger.info(f"File signature indicates ZIP format (possibly DOCX)")
            return ".docx"

        # Default to PDF as a fallback for most document processing
        logger.warning(
            f"Could not determine file type from signature, defaulting to PDF"
        )
        return ".pdf"
    except Exception as e:
        logger.error(f"Error detecting file type: {str(e)}")
        # Default to PDF as a fallback
        return ".pdf"


def fallback_extract_text(file_path: str) -> str:
    """
    Fallback method to extract text using specific libraries for each file type.

    Args:
        file_path (str): Path to the document file

    Returns:
        str: Extracted text
    """
    file_extension = os.path.splitext(file_path)[1].lower()

    # If no extension is detected, try to determine file type
    if not file_extension:
        file_extension = detect_file_type(file_path)
        logger.info(f"Using detected file type: '{file_extension}' for extraction")

    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
    elif file_extension == ".txt":
        return extract_text_from_txt(file_path)
    else:
        # For unknown file types, try PDF extraction as a fallback
        logger.warning(
            f"Unsupported file type: {file_extension}, attempting PDF extraction as fallback"
        )
        try:
            return extract_text_from_pdf(file_path)
        except Exception as pdf_error:
            logger.error(f"PDF extraction failed: {str(pdf_error)}")
            # As a last resort, try to read as plain text
            try:
                logger.warning("Attempting to read file as plain text")
                return extract_text_from_txt(file_path)
            except Exception as txt_error:
                logger.error(f"Plain text extraction failed: {str(txt_error)}")
                raise ValueError(
                    f"Could not extract text from file with extension: {file_extension}"
                )


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyPDF2.

    Args:
        file_path (str): Path to the PDF file

    Returns:
        str: Extracted text
    """
    try:
        # Try using PyPDF2 first
        import PyPDF2

        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

        # If PyPDF2 extracted meaningful text, return it
        if text.strip():
            return text

        # If PyPDF2 failed to extract meaningful text, try pdfplumber
        import pdfplumber

        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"

        return text

    except ImportError as e:
        logger.error(f"PDF extraction libraries not installed: {str(e)}")
        logger.error("Install with: pip install PyPDF2 pdfplumber")
        raise
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        file_path (str): Path to the DOCX file

    Returns:
        str: Extracted text
    """
    try:
        import docx

        doc = docx.Document(file_path)
        text = ""

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"

        return text
    except ImportError:
        logger.error(
            "python-docx is not installed. Install it with: pip install python-docx"
        )
        raise
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a TXT file.

    Args:
        file_path (str): Path to the TXT file

    Returns:
        str: Extracted text
    """
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        with open(file_path, "r", encoding="latin-1") as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.

    Args:
        text (str): Raw extracted text

    Returns:
        str: Cleaned text
    """
    import re

    # Remove excessive whitespace
    text = " ".join(text.split())

    # Replace multiple newlines with a single newline
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove non-printable characters
    text = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

    return text
